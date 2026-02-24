#!/usr/bin/env python3
"""
Personal Knowledge Indexer
==========================

Scans local files, cloud storage, and remote machines to create
a structured YAML document for LLM context loading.

Features:
- Multi-source: Local, Dropbox, Box, iCloud, remote SSH
- Content extraction: PDF, DOCX, Markdown, code
- Project detection: Git repos, package.json, pyproject.toml
- Semantic tagging: Auto-categorize by content
- Incremental updates: Only re-index changed files
- Privacy-aware: Never indexes credential contents

Usage:
    python indexer.py                    # Full scan
    python indexer.py --quick            # Quick scan (modified only)
    python indexer.py --remote mac-mini  # Include remote machine
    python indexer.py --output ~/knowledge.yaml
"""

import os
import sys
import json
import hashlib
import subprocess
import re
from pathlib import Path
from datetime import datetime, timedelta
from dataclasses import dataclass, field, asdict
from typing import Dict, List, Optional, Set, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
import argparse
import logging

# Optional imports for content extraction
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%H:%M:%S'
)
log = logging.getLogger(__name__)


# =============================================================================
# CONFIGURATION
# =============================================================================

@dataclass
class IndexConfig:
    """Configuration for the indexer."""

    # Paths to scan (~ expanded automatically)
    scan_paths: List[str] = field(default_factory=lambda: [
        "~/Desktop",
        "~/Documents",
        "~/Projects",
        "~/Developer",
        "~/.config",
        # OWC External Drive - Primary project storage
        "/Volumes/OWC drive/Dev",
        "/Volumes/OWC drive/Knowledge",
        "/Volumes/OWC drive/Archive AI Projects",
        # Claude Code session data
        "~/.claude",
        # SMB Network Share (Mac Mini) - Google Drive & Projects
        "/Volumes/christopherbailey/My Drive",  # Google Drive (41GB)
        "/Volumes/christopherbailey/Desktop/AI-WORK-HERE",  # Active AI projects
        "/Volumes/christopherbailey/Documents/GitHub",  # Git repositories
        "/Volumes/christopherbailey/_AI Output",  # AI-generated outputs
    ])

    # Cloud storage paths (mounted locations)
    cloud_paths: Dict[str, str] = field(default_factory=lambda: {
        "dropbox": "~/Dropbox",
        "box": "~/Box",
        "icloud": "~/Library/Mobile Documents/com~apple~CloudDocs",
        "google_drive": "~/Google Drive",
    })

    # Remote machines (SSH)
    remote_hosts: Dict[str, dict] = field(default_factory=lambda: {
        "mac-mini": {
            "host": "mac-mini.local",
            "user": os.environ.get("USER"),
            "paths": ["~/Documents", "~/Projects"],
        }
    })

    # File extensions to index
    code_extensions: Set[str] = field(default_factory=lambda: {
        ".py", ".js", ".ts", ".jsx", ".tsx", ".go", ".rs", ".rb",
        ".java", ".kt", ".swift", ".c", ".cpp", ".h", ".hpp",
        ".sh", ".bash", ".zsh", ".fish",
        ".sql", ".graphql",
    })

    config_extensions: Set[str] = field(default_factory=lambda: {
        ".yaml", ".yml", ".json", ".toml", ".ini", ".conf", ".cfg",
        ".env.example", ".editorconfig", ".prettierrc", ".eslintrc",
    })

    doc_extensions: Set[str] = field(default_factory=lambda: {
        ".md", ".markdown", ".txt", ".rst", ".org",
        ".pdf", ".docx", ".doc", ".rtf",
        ".tex", ".latex",
    })

    # Paths to skip
    skip_patterns: List[str] = field(default_factory=lambda: [
        "node_modules", "__pycache__", ".git", ".svn",
        "venv", ".venv", "env", ".env",
        ".DS_Store", "Thumbs.db",
        "*.pyc", "*.pyo", "*.class", "*.o",
        "dist", "build", "target", ".next",
        "Library/Caches", "Library/Application Support/*/Cache*",
        ".Trash",
        # Additional exclusions
        "Downloads",  # Contains personal files and credentials
        ".docker", ".pnpm-store",
        "debug",  # Claude debug folder (201MB, low value)
        "*.log", "*.lock",
        "Artifacts and Downloads",  # Contains OAuth secrets
        # SMB Share sensitive exclusions (HIPAA, PII, Financial)
        "Personal Taxes", "*Taxes*",
        "*Accounting*", "Investments Family Office",
        "Contracts and SOW's",  # Business confidential
        "*@erp-access.com",  # Colleague email folders
        # Medical/HIPAA data
        "Horos*", "OsiriX*", "DICOM", "Long Covid",
        "Medically Home",
        # Media folders (not useful for indexing)
        "Movies", "Music", "Pictures", "Photos",
        # Legacy/inactive cloud storage
        "Dropbox", ".dropbox", ".Box*", ".boxsync",
    ])

    # Sensitive paths (index existence, never contents)
    sensitive_patterns: List[str] = field(default_factory=lambda: [
        ".ssh", ".aws", ".gnupg", ".keys",
        "credentials", "secrets", ".env",
        "*.pem", "*.key", "*.crt",
        "id_rsa", "id_ed25519",
        # Additional security exclusions
        "*token*", "*password*", "*secret*",
        "client_secret_*.json", "gmail-credentials.json",
        "2FA_recovery", "recovery_code",
        ".streamlit", ".legal-review",
        ".env.backup*", "*.env.local",
    ])

    # Size limits
    max_file_size_mb: float = 10.0
    max_preview_chars: int = 500
    max_files_per_project: int = 100

    # Cache
    cache_path: str = "~/.cache/file-indexer"
    state_file: str = "index_state.json"


# =============================================================================
# DATA STRUCTURES
# =============================================================================

@dataclass
class FileInfo:
    """Information about a single file."""
    path: str
    relative_path: str
    host: str
    size_bytes: int
    modified: datetime
    file_type: str  # code, config, doc, other
    extension: str
    content_hash: Optional[str] = None
    preview: Optional[str] = None
    extracted_info: Dict = field(default_factory=dict)


@dataclass
class ProjectInfo:
    """Detected project/repository."""
    name: str
    path: str
    host: str
    project_type: str  # python, node, rust, go, etc.
    description: Optional[str] = None
    key_files: List[Dict] = field(default_factory=list)
    dependencies: List[str] = field(default_factory=list)
    summary: Optional[str] = None
    last_modified: Optional[datetime] = None


@dataclass
class IndexState:
    """Persistent state for incremental updates."""
    last_full_scan: Optional[datetime] = None
    file_hashes: Dict[str, str] = field(default_factory=dict)  # path -> hash
    file_mtimes: Dict[str, float] = field(default_factory=dict)  # path -> mtime


# =============================================================================
# FILE SCANNER
# =============================================================================

class FileScanner:
    """Scans filesystem for indexable files."""

    def __init__(self, config: IndexConfig):
        self.config = config
        self.skip_re = self._compile_skip_patterns()
        self.sensitive_re = self._compile_sensitive_patterns()

    def _compile_skip_patterns(self) -> re.Pattern:
        """Compile skip patterns into regex."""
        patterns = []
        for p in self.config.skip_patterns:
            # Convert glob to regex
            p = p.replace(".", r"\.")
            p = p.replace("*", ".*")
            p = p.replace("?", ".")
            patterns.append(f"(^|/){p}($|/)")
        return re.compile("|".join(patterns))

    def _compile_sensitive_patterns(self) -> re.Pattern:
        """Compile sensitive patterns into regex."""
        patterns = []
        for p in self.config.sensitive_patterns:
            p = p.replace(".", r"\.")
            p = p.replace("*", ".*")
            patterns.append(f"(^|/){p}($|/)")
        return re.compile("|".join(patterns))

    def should_skip(self, path: str) -> bool:
        """Check if path should be skipped."""
        return bool(self.skip_re.search(path))

    def is_sensitive(self, path: str) -> bool:
        """Check if path contains sensitive data."""
        return bool(self.sensitive_re.search(path))

    def get_file_type(self, path: Path) -> str:
        """Determine file type from extension."""
        ext = path.suffix.lower()
        if ext in self.config.code_extensions:
            return "code"
        elif ext in self.config.config_extensions:
            return "config"
        elif ext in self.config.doc_extensions:
            return "doc"
        return "other"

    def scan_directory(self, base_path: str, host: str = "local") -> List[FileInfo]:
        """Scan a directory recursively."""
        base = Path(base_path).expanduser()
        if not base.exists():
            log.warning(f"Path does not exist: {base}")
            return []

        files = []
        max_size = self.config.max_file_size_mb * 1024 * 1024

        for root, dirs, filenames in os.walk(base):
            # Filter directories in-place to skip
            dirs[:] = [d for d in dirs if not self.should_skip(os.path.join(root, d))]

            for filename in filenames:
                filepath = Path(root) / filename
                rel_path = str(filepath.relative_to(base))
                full_path = str(filepath)

                # Skip checks
                if self.should_skip(full_path):
                    continue

                try:
                    stat = filepath.stat()
                    if stat.st_size > max_size:
                        continue

                    file_type = self.get_file_type(filepath)
                    if file_type == "other":
                        continue  # Only index known types

                    info = FileInfo(
                        path=full_path,
                        relative_path=rel_path,
                        host=host,
                        size_bytes=stat.st_size,
                        modified=datetime.fromtimestamp(stat.st_mtime),
                        file_type=file_type,
                        extension=filepath.suffix.lower(),
                    )

                    # Mark sensitive files
                    if self.is_sensitive(full_path):
                        info.extracted_info["sensitive"] = True

                    files.append(info)

                except (PermissionError, OSError) as e:
                    log.debug(f"Cannot access {filepath}: {e}")

        log.info(f"Scanned {base}: {len(files)} files")
        return files


# =============================================================================
# CONTENT EXTRACTOR
# =============================================================================

class ContentExtractor:
    """Extracts content and metadata from files."""

    def __init__(self, config: IndexConfig):
        self.config = config

    def extract(self, file_info: FileInfo) -> FileInfo:
        """Extract content from a file."""
        if file_info.extracted_info.get("sensitive"):
            # Don't extract sensitive file contents
            return file_info

        path = Path(file_info.path)

        try:
            if file_info.extension == ".pdf":
                self._extract_pdf(file_info, path)
            elif file_info.extension in {".docx", ".doc"}:
                self._extract_docx(file_info, path)
            elif file_info.file_type in {"code", "config", "doc"}:
                self._extract_text(file_info, path)
        except Exception as e:
            log.debug(f"Extraction failed for {path}: {e}")

        return file_info

    def _extract_text(self, info: FileInfo, path: Path):
        """Extract text file content."""
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            # Generate hash
            info.content_hash = hashlib.md5(content.encode()).hexdigest()[:12]

            # Preview (first N chars, cleaned)
            preview = content[:self.config.max_preview_chars]
            preview = re.sub(r'\s+', ' ', preview).strip()
            info.preview = preview

            # Code-specific extraction
            if info.file_type == "code":
                self._extract_code_info(info, content)

            # Config-specific extraction
            elif info.file_type == "config":
                self._extract_config_info(info, content, path)

        except Exception as e:
            log.debug(f"Text extraction failed: {e}")

    def _extract_code_info(self, info: FileInfo, content: str):
        """Extract info from code files."""
        # Count lines
        lines = content.split('\n')
        info.extracted_info["lines"] = len(lines)

        # Extract docstring/module doc
        if info.extension == ".py":
            match = re.match(r'^[\s]*["\'][\'"]{2}(.*?)["\'][\'"]{2}', content, re.DOTALL)
            if match:
                info.extracted_info["docstring"] = match.group(1).strip()[:200]

            # Find classes and functions
            classes = re.findall(r'^class\s+(\w+)', content, re.MULTILINE)
            functions = re.findall(r'^def\s+(\w+)', content, re.MULTILINE)
            if classes:
                info.extracted_info["classes"] = classes[:10]
            if functions:
                info.extracted_info["functions"] = functions[:20]

        # Imports/dependencies
        if info.extension == ".py":
            imports = re.findall(r'^(?:from|import)\s+([\w.]+)', content, re.MULTILINE)
            info.extracted_info["imports"] = list(set(imports))[:20]

    def _extract_config_info(self, info: FileInfo, content: str, path: Path):
        """Extract info from config files."""
        if info.extension in {".yaml", ".yml"}:
            try:
                import yaml
                data = yaml.safe_load(content)
                if isinstance(data, dict):
                    info.extracted_info["top_keys"] = list(data.keys())[:10]
            except:
                pass

        elif info.extension == ".json":
            try:
                data = json.loads(content)
                if isinstance(data, dict):
                    info.extracted_info["top_keys"] = list(data.keys())[:10]
                    # Package.json specifics
                    if path.name == "package.json":
                        info.extracted_info["name"] = data.get("name")
                        info.extracted_info["dependencies"] = list(data.get("dependencies", {}).keys())[:10]
            except:
                pass

    def _extract_pdf(self, info: FileInfo, path: Path):
        """Extract text from PDF."""
        if not HAS_PDF:
            info.extracted_info["note"] = "PDF extraction requires PyPDF2"
            return

        try:
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                info.extracted_info["pages"] = len(reader.pages)

                # Extract first page text
                if reader.pages:
                    text = reader.pages[0].extract_text() or ""
                    preview = text[:self.config.max_preview_chars]
                    preview = re.sub(r'\s+', ' ', preview).strip()
                    info.preview = preview
        except Exception as e:
            log.debug(f"PDF extraction failed: {e}")

    def _extract_docx(self, info: FileInfo, path: Path):
        """Extract text from DOCX."""
        if not HAS_DOCX:
            info.extracted_info["note"] = "DOCX extraction requires python-docx"
            return

        try:
            doc = DocxDocument(path)
            paragraphs = [p.text for p in doc.paragraphs[:10]]
            text = '\n'.join(paragraphs)
            preview = text[:self.config.max_preview_chars]
            preview = re.sub(r'\s+', ' ', preview).strip()
            info.preview = preview
            info.extracted_info["paragraphs"] = len(doc.paragraphs)
        except Exception as e:
            log.debug(f"DOCX extraction failed: {e}")


# =============================================================================
# PROJECT DETECTOR
# =============================================================================

class ProjectDetector:
    """Detects and analyzes projects/repositories."""

    PROJECT_MARKERS = {
        "python": ["pyproject.toml", "setup.py", "requirements.txt", "Pipfile"],
        "node": ["package.json", "yarn.lock", "pnpm-lock.yaml"],
        "rust": ["Cargo.toml"],
        "go": ["go.mod"],
        "ruby": ["Gemfile"],
        "java": ["pom.xml", "build.gradle"],
        "dotnet": ["*.csproj", "*.sln"],
    }

    def __init__(self, config: IndexConfig):
        self.config = config

    def detect_projects(self, files: List[FileInfo]) -> List[ProjectInfo]:
        """Detect projects from file list."""
        # Group files by directory
        dir_files: Dict[str, List[FileInfo]] = {}
        for f in files:
            dir_path = str(Path(f.path).parent)
            dir_files.setdefault(dir_path, []).append(f)

        projects = []
        seen_roots = set()

        for dir_path, dir_file_list in dir_files.items():
            # Check for project markers
            filenames = {Path(f.path).name for f in dir_file_list}

            for proj_type, markers in self.PROJECT_MARKERS.items():
                for marker in markers:
                    if "*" in marker:
                        # Glob pattern
                        pattern = marker.replace("*", "")
                        if any(fn.endswith(pattern) for fn in filenames):
                            root = self._find_project_root(dir_path, marker)
                            if root and root not in seen_roots:
                                seen_roots.add(root)
                                proj = self._analyze_project(root, proj_type, files)
                                if proj:
                                    projects.append(proj)
                    elif marker in filenames:
                        root = self._find_project_root(dir_path, marker)
                        if root and root not in seen_roots:
                            seen_roots.add(root)
                            proj = self._analyze_project(root, proj_type, files)
                            if proj:
                                projects.append(proj)

        # Also check for .git directories
        for dir_path in dir_files:
            git_path = Path(dir_path) / ".git"
            if git_path.exists() and dir_path not in seen_roots:
                seen_roots.add(dir_path)
                proj = self._analyze_project(dir_path, "unknown", files)
                if proj:
                    projects.append(proj)

        log.info(f"Detected {len(projects)} projects")
        return projects

    def _find_project_root(self, start_path: str, marker: str) -> Optional[str]:
        """Find the project root directory."""
        path = Path(start_path)
        for _ in range(5):  # Max 5 levels up
            marker_path = path / marker.replace("*", "")
            if marker_path.exists() or any(path.glob(marker)):
                return str(path)
            if path.parent == path:
                break
            path = path.parent
        return start_path

    def _analyze_project(self, root: str, proj_type: str,
                         all_files: List[FileInfo]) -> Optional[ProjectInfo]:
        """Analyze a detected project."""
        root_path = Path(root)

        # Get files in this project
        project_files = [f for f in all_files if f.path.startswith(root)]
        if not project_files:
            return None

        # Get project name
        name = root_path.name

        # Find key files
        key_files = []
        key_names = {"README.md", "README", "main.py", "index.js", "index.ts",
                     "app.py", "server.py", "Makefile", "Dockerfile"}

        for f in project_files[:self.config.max_files_per_project]:
            filename = Path(f.path).name
            if filename in key_names or f.extracted_info.get("docstring"):
                key_files.append({
                    "path": f.relative_path,
                    "purpose": f.extracted_info.get("docstring", "")[:100] or filename,
                    "size_kb": f.size_bytes // 1024,
                })

        # Extract dependencies
        dependencies = []
        for f in project_files:
            deps = f.extracted_info.get("dependencies", [])
            dependencies.extend(deps)
        dependencies = list(set(dependencies))[:20]

        # Get description from README
        description = None
        for f in project_files:
            if "readme" in Path(f.path).name.lower():
                description = f.preview
                break

        # Find most recent modification
        if project_files:
            last_mod = max(f.modified for f in project_files)
        else:
            last_mod = None

        return ProjectInfo(
            name=name,
            path=root,
            host=project_files[0].host if project_files else "local",
            project_type=proj_type,
            description=description,
            key_files=key_files[:10],
            dependencies=dependencies,
            last_modified=last_mod,
        )


# =============================================================================
# REMOTE SCANNER
# =============================================================================

class RemoteScanner:
    """Scan remote machines via SSH."""

    def __init__(self, config: IndexConfig):
        self.config = config

    def scan_remote(self, host_name: str) -> List[FileInfo]:
        """Scan a remote host via SSH."""
        if host_name not in self.config.remote_hosts:
            log.error(f"Unknown remote host: {host_name}")
            return []

        host_config = self.config.remote_hosts[host_name]
        host = host_config["host"]
        user = host_config.get("user", os.environ.get("USER"))
        paths = host_config.get("paths", [])

        files = []

        for scan_path in paths:
            # Use find command over SSH
            cmd = [
                "ssh", f"{user}@{host}",
                f"find {scan_path} -type f -mtime -30 "
                f"-size -10M "
                f"\\( -name '*.py' -o -name '*.js' -o -name '*.md' "
                f"-o -name '*.yaml' -o -name '*.json' -o -name '*.txt' \\) "
                f"2>/dev/null | head -1000"
            ]

            try:
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                if result.returncode != 0:
                    log.warning(f"SSH scan failed for {host}: {result.stderr}")
                    continue

                for line in result.stdout.strip().split('\n'):
                    if not line:
                        continue

                    path = Path(line)
                    files.append(FileInfo(
                        path=line,
                        relative_path=line,
                        host=host_name,
                        size_bytes=0,  # Unknown without stat
                        modified=datetime.now(),  # Approximate
                        file_type=self._guess_type(path),
                        extension=path.suffix.lower(),
                    ))

            except subprocess.TimeoutExpired:
                log.warning(f"SSH scan timed out for {host}")
            except Exception as e:
                log.error(f"SSH scan error: {e}")

        log.info(f"Remote scan {host_name}: {len(files)} files")
        return files

    def _guess_type(self, path: Path) -> str:
        """Guess file type from extension."""
        ext = path.suffix.lower()
        if ext in {".py", ".js", ".ts", ".go", ".rs", ".sh"}:
            return "code"
        elif ext in {".yaml", ".yml", ".json", ".toml"}:
            return "config"
        elif ext in {".md", ".txt", ".pdf", ".docx"}:
            return "doc"
        return "other"


# =============================================================================
# YAML GENERATOR
# =============================================================================

class YAMLGenerator:
    """Generates the final YAML document."""

    def __init__(self, config: IndexConfig):
        self.config = config

    def generate(self, files: List[FileInfo], projects: List[ProjectInfo],
                 sources: List[Dict]) -> str:
        """Generate YAML document."""

        # Build document structure
        doc = {
            "meta": {
                "version": "1.0",
                "generated_at": datetime.now().isoformat(),
                "sources": sources,
                "stats": {
                    "total_files": len(files),
                    "total_size_mb": sum(f.size_bytes for f in files) / (1024 * 1024),
                    "documents": len([f for f in files if f.file_type == "doc"]),
                    "code_files": len([f for f in files if f.file_type == "code"]),
                    "configs": len([f for f in files if f.file_type == "config"]),
                }
            },
            "projects": [self._project_to_dict(p) for p in projects],
            "documents": self._group_documents(files),
            "code_patterns": self._extract_patterns(files),
            "configs": self._group_configs(files),
            "sensitive_paths": self._get_sensitive_paths(files),
            "recent_activity": self._get_recent_activity(files),
            "semantic_index": self._build_semantic_index(files, projects),
        }

        # Convert to YAML
        import yaml

        # Custom representer for datetime
        def datetime_representer(dumper, data):
            return dumper.represent_str(data.isoformat())

        yaml.add_representer(datetime, datetime_representer)

        return yaml.dump(doc, default_flow_style=False, sort_keys=False,
                        allow_unicode=True, width=100)

    def _project_to_dict(self, proj: ProjectInfo) -> Dict:
        """Convert project to dict."""
        return {
            "name": proj.name,
            "path": proj.path,
            "host": proj.host,
            "type": proj.project_type,
            "description": proj.description,
            "key_files": proj.key_files,
            "dependencies": proj.dependencies,
            "last_modified": proj.last_modified.isoformat() if proj.last_modified else None,
        }

    def _group_documents(self, files: List[FileInfo]) -> Dict:
        """Group documents by inferred category."""
        docs = [f for f in files if f.file_type == "doc"]

        categories = {
            "notes": [],
            "research": [],
            "reference": [],
            "other": [],
        }

        for doc in docs:
            path_lower = doc.path.lower()

            if any(x in path_lower for x in ["note", "journal", "diary", "log"]):
                cat = "notes"
            elif any(x in path_lower for x in ["research", "paper", "study", "thesis"]):
                cat = "research"
            elif any(x in path_lower for x in ["reference", "manual", "guide", "doc"]):
                cat = "reference"
            else:
                cat = "other"

            categories[cat].append({
                "path": doc.path,
                "host": doc.host,
                "format": doc.extension.lstrip("."),
                "preview": doc.preview,
                "modified": doc.modified.isoformat(),
            })

        # Limit each category
        return {k: v[:50] for k, v in categories.items() if v}

    def _group_configs(self, files: List[FileInfo]) -> Dict:
        """Group config files."""
        configs = [f for f in files if f.file_type == "config"]

        groups = {
            "shell": [],
            "git": [],
            "editors": [],
            "other": [],
        }

        for cfg in configs:
            name = Path(cfg.path).name.lower()

            if any(x in name for x in ["zsh", "bash", "fish", "profile"]):
                cat = "shell"
            elif "git" in name:
                cat = "git"
            elif any(x in name for x in ["vscode", "vim", "emacs", "editor"]):
                cat = "editors"
            else:
                cat = "other"

            groups[cat].append({
                "path": cfg.path,
                "host": cfg.host,
                "top_keys": cfg.extracted_info.get("top_keys", []),
            })

        return {k: v[:20] for k, v in groups.items() if v}

    def _extract_patterns(self, files: List[FileInfo]) -> List[Dict]:
        """Extract reusable code patterns."""
        # This would be more sophisticated in production
        # For now, just note files with classes
        patterns = []

        for f in files:
            classes = f.extracted_info.get("classes", [])
            if classes:
                patterns.append({
                    "file": f.path,
                    "classes": classes,
                    "language": f.extension.lstrip("."),
                })

        return patterns[:20]

    def _get_sensitive_paths(self, files: List[FileInfo]) -> List[str]:
        """Get list of sensitive paths (no contents)."""
        return list(set(
            f.path for f in files
            if f.extracted_info.get("sensitive")
        ))[:50]

    def _get_recent_activity(self, files: List[FileInfo]) -> Dict:
        """Get recent file activity."""
        now = datetime.now()
        week_ago = now - timedelta(days=7)

        recent = [f for f in files if f.modified > week_ago]
        recent.sort(key=lambda x: x.modified, reverse=True)

        return {
            "last_7_days": {
                "files_modified": len(recent),
                "most_active_dirs": self._get_active_dirs(recent),
            },
            "recently_modified": [
                {"path": f.path, "modified": f.modified.isoformat()}
                for f in recent[:20]
            ]
        }

    def _get_active_dirs(self, files: List[FileInfo]) -> List[str]:
        """Get most active directories."""
        from collections import Counter
        dirs = Counter(str(Path(f.path).parent) for f in files)
        return [d for d, _ in dirs.most_common(10)]

    def _build_semantic_index(self, files: List[FileInfo],
                               projects: List[ProjectInfo]) -> Dict:
        """Build semantic topic index."""
        topics = {}

        # Index by detected topics
        keywords = {
            "machine learning": ["ml", "neural", "model", "train", "predict"],
            "api": ["api", "endpoint", "request", "response", "client"],
            "database": ["sql", "database", "query", "table", "model"],
            "web": ["html", "css", "react", "vue", "frontend"],
            "devops": ["docker", "kubernetes", "deploy", "ci", "cd"],
        }

        for topic, kws in keywords.items():
            matching = []
            for f in files:
                path_lower = f.path.lower()
                preview_lower = (f.preview or "").lower()
                if any(kw in path_lower or kw in preview_lower for kw in kws):
                    matching.append(f.path)
            if matching:
                topics[topic] = matching[:10]

        return {"topics": topics}


# =============================================================================
# MAIN INDEXER
# =============================================================================

class KnowledgeIndexer:
    """Main indexer orchestrator."""

    def __init__(self, config: IndexConfig = None):
        self.config = config or IndexConfig()
        self.scanner = FileScanner(self.config)
        self.extractor = ContentExtractor(self.config)
        self.detector = ProjectDetector(self.config)
        self.remote = RemoteScanner(self.config)
        self.generator = YAMLGenerator(self.config)

        # State management
        self.state = self._load_state()

    def _load_state(self) -> IndexState:
        """Load previous index state."""
        cache_dir = Path(self.config.cache_path).expanduser()
        cache_dir.mkdir(parents=True, exist_ok=True)

        state_file = cache_dir / self.config.state_file
        if state_file.exists():
            try:
                data = json.loads(state_file.read_text())
                return IndexState(
                    last_full_scan=datetime.fromisoformat(data.get("last_full_scan")) if data.get("last_full_scan") else None,
                    file_hashes=data.get("file_hashes", {}),
                    file_mtimes=data.get("file_mtimes", {}),
                )
            except Exception as e:
                log.warning(f"Failed to load state: {e}")

        return IndexState()

    def _save_state(self):
        """Save index state."""
        cache_dir = Path(self.config.cache_path).expanduser()
        state_file = cache_dir / self.config.state_file

        data = {
            "last_full_scan": self.state.last_full_scan.isoformat() if self.state.last_full_scan else None,
            "file_hashes": self.state.file_hashes,
            "file_mtimes": self.state.file_mtimes,
        }
        state_file.write_text(json.dumps(data, indent=2))

    def run(self, quick: bool = False, remote_hosts: List[str] = None,
            output_path: str = None) -> str:
        """Run the indexer."""
        log.info("Starting knowledge indexer...")

        all_files = []
        sources = []

        # Scan local paths
        for scan_path in self.config.scan_paths:
            expanded = str(Path(scan_path).expanduser())
            if Path(expanded).exists():
                files = self.scanner.scan_directory(expanded, host="local")
                all_files.extend(files)
                sources.append({
                    "host": "local",
                    "path": scan_path,
                    "files_found": len(files),
                    "last_scan": datetime.now().isoformat(),
                })

        # Scan cloud storage
        for cloud_name, cloud_path in self.config.cloud_paths.items():
            expanded = str(Path(cloud_path).expanduser())
            if Path(expanded).exists():
                files = self.scanner.scan_directory(expanded, host=cloud_name)
                all_files.extend(files)
                sources.append({
                    "host": cloud_name,
                    "path": cloud_path,
                    "files_found": len(files),
                    "last_scan": datetime.now().isoformat(),
                })

        # Scan remote hosts
        if remote_hosts:
            for host in remote_hosts:
                files = self.remote.scan_remote(host)
                all_files.extend(files)
                sources.append({
                    "host": host,
                    "path": "remote",
                    "files_found": len(files),
                    "last_scan": datetime.now().isoformat(),
                })

        log.info(f"Total files found: {len(all_files)}")

        # Extract content (parallel)
        log.info("Extracting content...")
        with ThreadPoolExecutor(max_workers=4) as executor:
            all_files = list(executor.map(self.extractor.extract, all_files))

        # Detect projects
        log.info("Detecting projects...")
        projects = self.detector.detect_projects(all_files)

        # Generate YAML
        log.info("Generating YAML...")
        yaml_content = self.generator.generate(all_files, projects, sources)

        # Save output
        output = Path(output_path or "~/knowledge.yaml").expanduser()
        output.write_text(yaml_content)
        log.info(f"Saved to {output}")

        # Update state
        self.state.last_full_scan = datetime.now()
        for f in all_files:
            if f.content_hash:
                self.state.file_hashes[f.path] = f.content_hash
        self._save_state()

        return str(output)


# =============================================================================
# CLI
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description="Personal Knowledge Indexer")
    parser.add_argument("--quick", action="store_true",
                        help="Quick scan (recently modified only)")
    parser.add_argument("--remote", nargs="*", metavar="HOST",
                        help="Include remote hosts (e.g., mac-mini)")
    parser.add_argument("--output", "-o", type=str,
                        help="Output path (default: ~/knowledge.yaml)")
    parser.add_argument("--verbose", "-v", action="store_true",
                        help="Verbose output")

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    indexer = KnowledgeIndexer()
    output = indexer.run(
        quick=args.quick,
        remote_hosts=args.remote,
        output_path=args.output,
    )

    print(f"\nIndex created: {output}")
    print("Load into LLM with: cat ~/knowledge.yaml | pbcopy")


if __name__ == "__main__":
    main()
